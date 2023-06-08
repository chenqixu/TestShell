# coding=utf-8
#############################################################################################
# 程序描述：excel转成word
# 实现功能：
# 运行周期：
# 创建作者：
# 创建日期：2023-06-07
# 特别注意：
#############################################################################################
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

import sys
reload(sys)
sys.setdefaultencoding('utf-8')
print(sys.getdefaultencoding())

# 新建文档
d = Document()

# 第一种设置标题的方式
d.add_heading("标题0".encode().decode()+"one hello world", level=0)

# 第二种设置标题的方式，此方式还可以设置文本的字体、颜色、大小等属性
run = d.add_heading("", level=1).add_run("标题1".encode().decode())
# 设置西文字体
run.font.name = u'宋体'
# 设置中文字体
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
# 设置字体颜色
run.font.color.rgb = RGBColor(255, 55, 55)  # 红色
# 设置字体大小
run.font.size = Pt(30)
# 设置下划线
run.font.underline = True
# 设置删除线
run.font.strike = True
# 设置加粗
run.bold = True
# 设置斜体
run.italic = True
# 保存文档
d.save("d:/tmp/data/word/test.docx")